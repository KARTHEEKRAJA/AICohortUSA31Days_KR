import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from PIL import Image, ImageDraw, ImageFont

os.makedirs("data1/raw", exist_ok=True)

# --- 1. SBC-style PDF ---
c = canvas.Canvas("data1/raw/sbc_gold_ppo.pdf", pagesize=letter)
c.setFont("Helvetica-Bold", 16)
c.drawString(72, 740, "Summary of Benefits and Coverage: Gold PPO (P101)")
c.setFont("Helvetica", 11)
lines = [
    "Coverage Period: 01/01/2026 - 12/31/2026",
    "Plan Type: PPO | Network Tier: Gold",
    "",
    "Overall deductible: $2,000 individual / $4,000 family",
    "Out-of-pocket limit: $6,500 individual / $13,000 family",
    "Primary care visit: $25 copay per visit",
    "Specialist visit: $50 copay per visit",
    "Emergency room care: 10% coinsurance after deductible",
    "Generic drugs: $10 copay | Preferred brand drugs: $40 copay",
    "X-rays and diagnostic imaging: 10% coinsurance",
    "Outpatient surgery: 10% coinsurance after deductible",
    "",
    "Services NOT covered: cosmetic surgery, weight loss programs,",
    "long-term care, non-emergency care when traveling outside the U.S.",
]
y = 710
for ln in lines:
    c.drawString(72, y, ln)
    y -= 20
c.save()

# --- 2. Claims-process Word doc ---
doc = Document()
doc.add_heading("How to File a Claim - Member Guide", 0)
doc.add_heading("Step 1: Collect Your Documents", level=1)
doc.add_paragraph("Obtain an itemized bill from your provider showing the date of "
                  "service, procedure codes, and charges.")
doc.add_heading("Step 2: Complete the Claim Form", level=1)
doc.add_paragraph("Fill out form CF-100 with your member ID, plan ID, and provider details. "
                  "Claims must be submitted within 90 days of the date of service.")
doc.add_heading("Step 3: Submit", level=1)
doc.add_paragraph("Mail to: Claims Department, PO Box 12345, Hartford, CT, or upload "
                  "via the member portal. Processing takes 15-30 business days.")
doc.add_heading("Appeals", level=1)
doc.add_paragraph("If a claim is denied, you may appeal within 180 days. Include the "
                  "denial letter and any supporting medical records.")
doc.save("data1/raw/claims_process.docx")

# --- 3. "Scanned" enrollment form image ---
img = Image.new("RGB", (1000, 1200), "white")
d = ImageDraw.Draw(img)
try:
    font_b = ImageFont.truetype("arial.ttf", 34)
    font = ImageFont.truetype("arial.ttf", 26)
except OSError:
    font_b = font = ImageFont.load_default()

d.text((60, 50), "HEALTH PLAN ENROLLMENT FORM", font=font_b, fill="black")
fields = [
    "Member Name: Jane Q. Sample", "Date of Birth: 04/12/1990",
    "Member ID: M2001", "Selected Plan: Silver HMO (P102)",
    "Coverage Start Date: 08/01/2026", "Dependents: 1 (child)",
    "Primary Care Physician: Dr. A. Rivera",
    "Signature: Jane Q. Sample     Date: 07/20/2026",
]
y = 150
for f in fields:
    d.text((60, y), f, font=font, fill="black")
    d.line((60, y + 40, 940, y + 40), fill="gray", width=2)
    y += 110
img.save("data1/raw/enrollment_form_scan.png")

print("Created 3 docs in data1/raw/")